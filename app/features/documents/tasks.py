# """
# Background tasks for document processing.

# Tasks run in Celery workers, separate from the API server.
# """

# import logging
# import traceback
# from datetime import datetime, timedelta

# from celery import Task as CeleryTask
# from sqlalchemy import select

# from app.core.celery_app import celery_app
# from app.core.database import db_manager
# from app.models.document import Document, DocumentStatus
# from app.models.task import Task, TaskStatus

# logger = logging.getLogger(__name__)


# class DatabaseTask(CeleryTask):
#     """
#     Base task class with database session management.
    
#     Provides automatic database session handling for tasks.
#     """
    
#     _db = None
    
#     def after_return(self, *args, **kwargs):
#         """Close database connection after task completes."""
#         if self._db is not None:
#             self._db.close()


# @celery_app.task(bind=True, base=DatabaseTask, name="process_document")
# def process_document(self, document_id: str, tenant_id: str) -> dict:
#     """
#     Main document processing task.
    
#     Orchestrates the document processing pipeline:
#     1. Extract text content
#     2. Count pages (for PDFs)
#     3. Generate metadata
#     4. Update document status
    
#     Args:
#         document_id: Document ID to process
#         tenant_id: Tenant ID for isolation
        
#     Returns:
#         Processing result dictionary
#     """
#     logger.info(f"Starting document processing: {document_id}")
    
#     # Create task tracking record
#     db_manager.init()
    
#     import asyncio
#     loop = asyncio.new_event_loop()
#     asyncio.set_event_loop(loop)
    
#     try:
#         result = loop.run_until_complete(
#             _process_document_async(self, document_id, tenant_id)
#         )
#         return result
#     finally:
#         loop.close()


# async def _process_document_async(task_instance, document_id: str, tenant_id: str) -> dict:
#     """Async implementation of document processing."""
    
#     async for db in db_manager.get_session():
#         try:
#             # Create task tracking
#             task_record = Task(
#                 task_id=task_instance.request.id,
#                 task_name="process_document",
#                 task_type="document_processing",
#                 status=TaskStatus.STARTED,
#                 started_at=datetime.utcnow(),
#                 resource_type="document",
#                 resource_id=document_id,
#                 tenant_id=tenant_id,
#                 progress=0,
#             )
#             db.add(task_record)
#             await db.commit()
            
#             # Fetch document
#             result = await db.execute(
#                 select(Document).where(Document.id == document_id)
#             )
#             document = result.scalar_one_or_none()
            
#             if not document:
#                 raise ValueError(f"Document not found: {document_id}")
            
#             # Update document status
#             document.status = DocumentStatus.PROCESSING
#             document.processing_started_at = datetime.utcnow()
#             await db.commit()
            
#             # Update progress
#             task_record.progress = 10
#             await db.commit()
            
#             # Step 1: Extract text
#             logger.info(f"Extracting text from {document.filename}")
#             text_content = await extract_text_from_file(document)
            
#             task_record.progress = 50
#             await db.commit()
            
#             # Step 2: Count pages (for PDFs)
#             page_count = None
#             if document.document_type == "pdf":
#                 page_count = await count_pdf_pages(document)
            
#             task_record.progress = 70
#             await db.commit()
            
#             # Step 3: Update document with extracted data
#             document.text_content = text_content[:10000] if text_content else None  # Limit to 10k chars
#             document.page_count = page_count
#             document.status = DocumentStatus.COMPLETED
#             document.processing_completed_at = datetime.utcnow()
            
#             task_record.progress = 90
#             await db.commit()
            
#             # Complete task
#             task_record.status = TaskStatus.SUCCESS
#             task_record.progress = 100
#             task_record.completed_at = datetime.utcnow()
#             task_record.result = {
#                 "text_length": len(text_content) if text_content else 0,
#                 "page_count": page_count,
#                 "processing_time_seconds": (
#                     document.processing_completed_at - document.processing_started_at
#                 ).total_seconds(),
#             }
#             await db.commit()
            
#             logger.info(f"Document processing completed: {document_id}")
            
#             return {
#                 "document_id": document_id,
#                 "status": "completed",
#                 "text_length": len(text_content) if text_content else 0,
#                 "page_count": page_count,
#             }
            
#         except Exception as e:
#             logger.error(f"Document processing failed: {document_id} - {e}")
            
#             # Update document status
#             if document:
#                 document.status = DocumentStatus.FAILED
#                 document.error_message = str(e)
#                 await db.commit()
            
#             # Update task status
#             task_record.status = TaskStatus.FAILURE
#             task_record.error = str(e)
#             task_record.traceback = traceback.format_exc()
#             task_record.completed_at = datetime.utcnow()
#             await db.commit()
            
#             # Retry logic
#             if task_record.retry_count < task_record.max_retries:
#                 task_record.retry_count += 1
#                 task_record.status = TaskStatus.RETRY
#                 await db.commit()
                
#                 # Retry after delay
#                 raise task_instance.retry(exc=e, countdown=60 * (2 ** task_record.retry_count))
            
#             raise


# async def extract_text_from_file(document: Document) -> str | None:
#     """
#     Extract text content from document file.
    
#     Supports:
#     - PDF: PyPDF2
#     - Word: python-docx
#     - Text/Markdown: Direct read
#     """
#     from app.features.documents.storage import get_storage
    
#     storage = get_storage()
    
#     try:
#         # Get file content
#         file_content = await storage.get(document.file_path)
        
#         # Extract based on document type
#         if document.document_type == "pdf":
#             return extract_text_from_pdf(file_content)
#         elif document.document_type == "word":
#             return extract_text_from_docx(file_content)
#         elif document.document_type in ["text", "markdown"]:
#             return file_content.decode("utf-8")
#         else:
#             logger.warning(f"Unsupported document type for text extraction: {document.document_type}")
#             return None
            
#     except Exception as e:
#         logger.error(f"Text extraction failed: {e}")
#         return None


# def extract_text_from_pdf(file_content: bytes) -> str:
#     """Extract text from PDF using PyPDF2."""
#     import PyPDF2
#     from io import BytesIO
    
#     text_parts = []
    
#     try:
#         pdf_file = BytesIO(file_content)
#         pdf_reader = PyPDF2.PdfReader(pdf_file)
        
#         for page in pdf_reader.pages:
#             text = page.extract_text()
#             if text:
#                 text_parts.append(text)
        
#         return "\n\n".join(text_parts)
        
#     except Exception as e:
#         logger.error(f"PDF text extraction error: {e}")
#         return ""


# def extract_text_from_docx(file_content: bytes) -> str:
#     """Extract text from Word document using python-docx."""
#     import docx
#     from io import BytesIO
    
#     try:
#         doc_file = BytesIO(file_content)
#         doc = docx.Document(doc_file)
        
#         text_parts = []
#         for paragraph in doc.paragraphs:
#             if paragraph.text.strip():
#                 text_parts.append(paragraph.text)
        
#         return "\n\n".join(text_parts)
        
#     except Exception as e:
#         logger.error(f"DOCX text extraction error: {e}")
#         return ""


# async def count_pdf_pages(document: Document) -> int | None:
#     """Count pages in PDF document."""
#     import PyPDF2
#     from io import BytesIO
#     from app.features.documents.storage import get_storage
    
#     storage = get_storage()
    
#     try:
#         file_content = await storage.get(document.file_path)
#         pdf_file = BytesIO(file_content)
#         pdf_reader = PyPDF2.PdfReader(pdf_file)
#         return len(pdf_reader.pages)
        
#     except Exception as e:
#         logger.error(f"PDF page count error: {e}")
#         return None


# @celery_app.task(name="cleanup_failed_documents")
# def cleanup_failed_documents() -> dict:
#     """
#     Periodic task to retry or cleanup failed documents.
    
#     Runs every hour via Celery Beat.
#     """
#     logger.info("Starting failed document cleanup")
    
#     db_manager.init()
#     import asyncio
#     loop = asyncio.new_event_loop()
#     asyncio.set_event_loop(loop)
    
#     try:
#         result = loop.run_until_complete(_cleanup_failed_documents_async())
#         return result
#     finally:
#         loop.close()


# async def _cleanup_failed_documents_async() -> dict:
#     """Async implementation of cleanup."""
    
#     async for db in db_manager.get_session():
#         # Find documents stuck in processing (>1 hour)
#         one_hour_ago = datetime.utcnow() - timedelta(hours=1)
        
#         result = await db.execute(
#             select(Document).where(
#                 Document.status == DocumentStatus.PROCESSING,
#                 Document.processing_started_at < one_hour_ago
#             )
#         )
        
#         stuck_documents = result.scalars().all()
        
#         for doc in stuck_documents:
#             logger.warning(f"Resetting stuck document: {doc.id}")
#             doc.status = DocumentStatus.FAILED
#             doc.error_message = "Processing timeout - exceeded 1 hour"
        
#         await db.commit()
        
#         logger.info(f"Cleanup complete: {len(stuck_documents)} documents reset")
        
#         return {
#             "documents_reset": len(stuck_documents),
#         }


# @celery_app.task(name="generate_daily_stats")
# def generate_daily_stats() -> dict:
#     """
#     Generate daily statistics report.
    
#     This is an example periodic task that could:
#     - Calculate daily upload counts
#     - Generate tenant usage reports
#     - Send summary emails to admins
#     """
#     logger.info("Generating daily statistics")
    
#     # TODO: Implement actual statistics generation
    
#     return {
#         "status": "completed",
#         "timestamp": datetime.utcnow().isoformat(),
#     }



"""
Background tasks for document processing.
Tasks run in Celery workers, separate from the API server.
"""

import logging
import traceback
import asyncio
from datetime import datetime, timedelta, timezone

from celery import Task as CeleryTask
from sqlalchemy import select

from app.core.celery_app import celery_app
from app.core.database import db_manager
from app.models.document import Document, DocumentStatus
from app.models.task import Task, TaskStatus

logger = logging.getLogger(__name__)

class DatabaseTask(CeleryTask):
    """Base task class with database session management."""
    _db = None
    
    def after_return(self, *args, **kwargs):
        """Close database connection after task completes."""
        if self._db is not None:
            self._db.close()

def get_or_create_event_loop():
    """Helper to handle asyncio loops safely within thread-based workers."""
    try:
        return asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        return loop

@celery_app.task(bind=True, base=DatabaseTask, name="process_document")
def process_document(self, document_id: str, tenant_id: str) -> dict:
    """Orchestrates the document processing pipeline."""
    logger.info(f"Starting document processing: {document_id}")
    
    db_manager.init()
    loop = get_or_create_event_loop()
    
    try:
        return loop.run_until_complete(
            _process_document_async(self, document_id, tenant_id)
        )
    except Exception as exc:
        # This catch handles the retry logic if the async part fails
        logger.error(f"Task wrapper caught exception: {exc}")
        raise exc

async def _process_document_async(task_instance, document_id: str, tenant_id: str) -> dict:
    """Async implementation of document processing."""
    # PREVENT UnboundLocalError: Initialize variables at the start
    document = None
    task_record = None
    
    async for db in db_manager.get_session():
        try:
            # 1. Create task tracking record
            task_record = Task(
                task_id=task_instance.request.id,
                task_name="process_document",
                task_type="document_processing",
                status=TaskStatus.STARTED,
                started_at=datetime.now(timezone.utc),
                resource_type="document",
                resource_id=document_id,
                tenant_id=tenant_id,
                progress=0,
            )
            db.add(task_record)
            await db.commit()
            
            # 2. Fetch document
            result = await db.execute(
                select(Document).where(Document.id == document_id)
            )
            document = result.scalar_one_or_none()
            
            if not document:
                raise ValueError(f"Document not found: {document_id}")
            
            # 3. Update document status to processing
            document.status = DocumentStatus.PROCESSING
            document.processing_started_at = datetime.now(timezone.utc)
            task_record.progress = 10
            await db.commit()
            
            # 4. Step 1: Extract text
            logger.info(f"Extracting text from {document.filename}")
            text_content = await extract_text_from_file(document)
            
            task_record.progress = 50
            await db.commit()
            
            # 5. Step 2: Count pages (for PDFs)
            page_count = None
            if document.document_type == "pdf":
                page_count = await count_pdf_pages(document)
            
            task_record.progress = 70
            await db.commit()
            
            # 6. Step 3: Final Update
            document.text_content = text_content[:10000] if text_content else None
            document.page_count = page_count
            document.status = DocumentStatus.COMPLETED
            document.processing_completed_at = datetime.now(timezone.utc)
            
            task_record.status = TaskStatus.SUCCESS
            task_record.progress = 100
            task_record.completed_at = datetime.now(timezone.utc)
            task_record.result = {
                "text_length": len(text_content) if text_content else 0,
                "page_count": page_count,
                "processing_time_seconds": (
                    document.processing_completed_at - document.processing_started_at
                ).total_seconds(),
            }
            await db.commit()
            
            logger.info(f"Document processing completed: {document_id}")
            return {"document_id": document_id, "status": "completed"}
            
        except Exception as e:
            logger.error(f"Document processing failed: {document_id} - {e}")
            
            # Safely update document if it was successfully fetched
            if document:
                document.status = DocumentStatus.FAILED
                document.error_message = str(e)
            
            # Safely update task tracking if it was created
            if task_record:
                task_record.status = TaskStatus.FAILURE
                task_record.error = str(e)
                task_record.traceback = traceback.format_exc()
                task_record.completed_at = datetime.now(timezone.utc)
                
                # Handle Retries
                if task_record.retry_count < task_record.max_retries:
                    task_record.retry_count += 1
                    task_record.status = TaskStatus.RETRY
                    await db.commit()
                    countdown = 60 * (2 ** task_record.retry_count)
                    raise task_instance.retry(exc=e, countdown=countdown)
            
            await db.commit()
            raise e

async def extract_text_from_file(document: Document) -> str | None:
    """Extract text content from document file via Storage."""
    from app.features.documents.storage import get_storage
    storage = get_storage()
    
    try:
        file_content = await storage.get(document.file_path)
        if not file_content:
            return None

        if document.document_type == "pdf":
            return extract_text_from_pdf(file_content)
        elif document.document_type == "word":
            return extract_text_from_docx(file_content)
        elif document.document_type in ["text", "markdown"]:
            return file_content.decode("utf-8")
        return None
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return None

def extract_text_from_pdf(file_content: bytes) -> str:
    import PyPDF2
    from io import BytesIO
    text_parts = []
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF error: {e}")
        return ""

def extract_text_from_docx(file_content: bytes) -> str:
    import docx
    from io import BytesIO
    try:
        doc = docx.Document(BytesIO(file_content))
        return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        logger.error(f"DOCX error: {e}")
        return ""

async def count_pdf_pages(document: Document) -> int | None:
    import PyPDF2
    from io import BytesIO
    from app.features.documents.storage import get_storage
    storage = get_storage()
    try:
        file_content = await storage.get(document.file_path)
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        return len(pdf_reader.pages)
    except Exception:
        return None

@celery_app.task(name="cleanup_failed_documents")
def cleanup_failed_documents() -> dict:
    """Periodic task to cleanup stuck processing jobs."""
    logger.info("Starting failed document cleanup")
    db_manager.init()
    loop = get_or_create_event_loop()
    return loop.run_until_complete(_cleanup_failed_documents_async())

async def _cleanup_failed_documents_async() -> dict:
    async for db in db_manager.get_session():
        one_hour_ago = datetime.now(timezone.utc) - timedelta(hours=1)
        result = await db.execute(
            select(Document).where(
                Document.status == DocumentStatus.PROCESSING,
                Document.processing_started_at < one_hour_ago
            )
        )
        stuck_docs = result.scalars().all()
        for doc in stuck_docs:
            doc.status = DocumentStatus.FAILED
            doc.error_message = "Processing timeout"
        await db.commit()
        return {"documents_reset": len(stuck_docs)}

@celery_app.task(name="generate_daily_stats")
def generate_daily_stats() -> dict:
    logger.info("Generating daily statistics")
    return {"status": "completed", "timestamp": datetime.now(timezone.utc).isoformat()}